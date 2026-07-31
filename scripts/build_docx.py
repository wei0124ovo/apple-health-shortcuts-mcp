"""Build the shareable Chinese tutorial DOCX from GUIDE.zh-CN.md."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "GUIDE-MOBILE.zh-CN.md"
OUTPUT_DIR = ROOT / "dist"
OUTPUT = OUTPUT_DIR / "Apple-Health-快捷指令-MCP-从零搭建教程.docx"

INK = "172033"
MUTED = "526174"
BLUE = "2563EB"
TEAL = "0F766E"
PALE_BLUE = "EAF2FF"
PALE_TEAL = "E8F7F4"
PALE_GRAY = "F3F5F8"
WHITE = "FFFFFF"
BORDER = "CCD5E1"
CODE_BG = "101827"
CODE_FG = "F8FAFC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    relation_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.extend([color, underline])
    run.append(props)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


INLINE = re.compile(
    r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\(https?://[^)]+\))"
)


def add_inline(paragraph, text: str, *, color: str = INK) -> None:
    cursor = 0
    for match in INLINE.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            run.font.color.rgb = RGBColor.from_string(color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(color)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(TEAL)
        else:
            link_match = re.match(r"\[([^\]]+)\]\((https?://[^)]+)\)", token)
            if link_match:
                add_hyperlink(paragraph, link_match.group(1), link_match.group(2))
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    number = OxmlElement("w:t")
    number.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, number, end])
    paragraph.add_run(" 页")


def set_east_asian_font(style, name: str) -> None:
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    set_east_asian_font(normal, "Microsoft YaHei")
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(4.5)
    normal.paragraph_format.line_spacing = 1.16

    for name, size, color, before, after in (
        ("Title", 28, INK, 0, 12),
        ("Subtitle", 13, MUTED, 0, 8),
        ("Heading 1", 19, BLUE, 18, 8),
        ("Heading 2", 14, TEAL, 13, 6),
        ("Heading 3", 11.5, INK, 10, 4),
    ):
        style = styles[name]
        set_east_asian_font(style, "Microsoft YaHei")
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.left_indent = Mm(0)
        style.paragraph_format.first_line_indent = Mm(0)

    if "Code Block" not in styles:
        code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = styles["Code Block"]
    set_east_asian_font(code_style, "Microsoft YaHei")
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(8.5)
    code_style.font.color.rgb = RGBColor.from_string(CODE_FG)
    code_style.paragraph_format.left_indent = Mm(4)
    code_style.paragraph_format.right_indent = Mm(4)
    code_style.paragraph_format.space_before = Pt(3)
    code_style.paragraph_format.space_after = Pt(6)
    code_style.paragraph_format.keep_together = True

    if "Callout" not in styles:
        callout = styles.add_style("Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = styles["Callout"]
    set_east_asian_font(callout, "Microsoft YaHei")
    callout.font.size = Pt(9.5)
    callout.font.color.rgb = RGBColor.from_string(TEAL)
    callout.paragraph_format.left_indent = Mm(6)
    callout.paragraph_format.right_indent = Mm(4)
    callout.paragraph_format.space_before = Pt(4)
    callout.paragraph_format.space_after = Pt(7)
    callout.paragraph_format.keep_together = True


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)


def add_cover(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Mm(23)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)

    label = document.add_paragraph()
    label.paragraph_format.space_after = Pt(54)
    label_run = label.add_run("独立开源 · 每人一套 · AI 只读")
    label_run.bold = True
    label_run.font.size = Pt(10)
    label_run.font.color.rgb = RGBColor.from_string(TEAL)

    title = document.add_paragraph(style="Title")
    title.add_run("让 AI 看懂我的\nApple 健康")

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.add_run("中国时区默认版：只用手机 + Cloudflare 一键部署 + HTTP MCP")

    rule = document.add_paragraph()
    rule.paragraph_format.space_before = Pt(14)
    rule.paragraph_format.space_after = Pt(18)
    shade_paragraph(rule, BLUE)
    rule.add_run(" ")

    intro = document.add_paragraph()
    intro.paragraph_format.space_after = Pt(10)
    run = intro.add_run(
        "不用电脑，不改代码，不填数据库编号。"
        "手机一键部署后，网页自动生成钥匙。"
    )
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string(INK)

    box = document.add_table(rows=3, cols=2)
    box.alignment = WD_TABLE_ALIGNMENT.LEFT
    box.autofit = False
    for row in box.rows:
        row.cells[0].width = Cm(4.0)
        row.cells[1].width = Cm(11.8)
        for cell in row.cells:
            set_cell_shading(cell, PALE_BLUE)
            set_cell_border(cell, PALE_BLUE, "0")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    facts = [
        ("默认时区", "中国 Asia/Shanghai"),
        ("所需设备", "一台 iPhone；Apple Watch 可选"),
        ("AI 权限", "只有 3 个只读工具"),
    ]
    for row, (left, right) in zip(box.rows, facts):
        p1 = row.cells[0].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(left)
        r1.bold = True
        r1.font.color.rgb = RGBColor.from_string(BLUE)
        p2 = row.cells[1].paragraphs[0]
        p2.paragraph_format.space_after = Pt(0)
        p2.add_run(right)

    document.add_paragraph()
    version = document.add_paragraph()
    version.paragraph_format.space_before = Pt(54)
    vrun = version.add_run("手机分享版教程 · 2.0 · 2026 年 7 月")
    vrun.font.size = Pt(9)
    vrun.font.color.rgb = RGBColor.from_string(MUTED)
    document.add_page_break()


def add_running_header_footer(section) -> None:
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("APPLE HEALTH 快捷指令 MCP  ·  从零搭建教程")
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(BLUE)
    shade_paragraph(p, PALE_BLUE)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.paragraph_format.space_before = Pt(2)
    add_page_number(fp)
    for run in fp.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string(MUTED)


def add_contents(document: Document) -> None:
    document.add_heading("拿着这张地图走", level=1)
    p = document.add_paragraph()
    add_inline(p, "先在手机一键部署并领取钥匙，再安装快捷指令，最后把 MCP 交给 AI。")
    items = [
        "第一～二部分　手机一键部署，打开设置页领取钥匙",
        "第三～五部分　安装或制作健康、睡眠快捷指令",
        "第六～七部分　自动运行并把 3 个工具交给 AI",
        "第八部分　群主制作 iCloud 链接，让群友点链接安装",
        "最后　优缺点与完成检查",
    ]
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        add_inline(paragraph, item)

    note = document.add_paragraph(style="Callout")
    shade_paragraph(note, PALE_TEAL)
    add_inline(
        note,
        "重要：公开分享代码和教程没有问题，但任何人的 Worker 地址、部署密码和两把钥匙都不能跟着公开。",
        color=TEAL,
    )


def split_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def is_separator_row(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    columns = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, (docx_row, values) in enumerate(zip(table.rows, rows)):
        prevent_row_split(docx_row)
        if row_index == 0:
            set_repeat_table_header(docx_row)
        for column_index, cell in enumerate(docx_row.cells):
            value = values[column_index] if column_index < len(values) else ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            set_cell_shading(cell, BLUE if row_index == 0 else (PALE_GRAY if row_index % 2 == 0 else WHITE))
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(1.5)
            paragraph.paragraph_format.space_after = Pt(1.5)
            if len(rows) <= 6 and row_index < len(rows) - 1:
                paragraph.paragraph_format.keep_with_next = True
            add_inline(paragraph, value, color=WHITE if row_index == 0 else INK)
            for run in paragraph.runs:
                run.font.size = Pt(8.3 if columns >= 4 else 9)
                if row_index == 0:
                    run.bold = True
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code(document: Document, code_lines: list[str]) -> None:
    paragraph = document.add_paragraph(style="Code Block")
    shade_paragraph(paragraph, CODE_BG)
    run = paragraph.add_run("\n".join(code_lines))
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(CODE_FG)


def add_markdown_body(document: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    index = 0
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if not paragraph_buffer:
            return
        text = " ".join(piece.strip() for piece in paragraph_buffer).strip()
        if text:
            paragraph = document.add_paragraph()
            if text.endswith(("：", ":")):
                paragraph.paragraph_format.keep_with_next = True
            add_inline(paragraph, text)
        paragraph_buffer = []

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("# ") or stripped.startswith("## 完全免费"):
            index += 1
            continue

        if stripped.startswith("```"):
            flush_paragraph()
            index += 1
            code_lines: list[str] = []
            while index < len(lines) and not lines[index].strip().startswith("```"):
                code_lines.append(lines[index])
                index += 1
            add_code(document, code_lines)
            index += 1
            continue

        if (
            stripped.startswith("|")
            and index + 1 < len(lines)
            and is_separator_row(lines[index + 1])
        ):
            flush_paragraph()
            table_rows = [split_table_row(line)]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_rows.append(split_table_row(lines[index]))
                index += 1
            add_table(document, table_rows)
            continue

        heading = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading:
            flush_paragraph()
            level = min(3, len(heading.group(1)) - 1)
            heading_text = heading.group(2).replace("`", "")
            if level == 1:
                heading_text = "\u2009" + heading_text
            document.add_heading(heading_text, level=level)
            index += 1
            continue

        if stripped.startswith(">"):
            flush_paragraph()
            quote_lines = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                quote_lines.append(lines[index].strip().lstrip(">").strip())
                index += 1
            paragraph = document.add_paragraph(style="Callout")
            shade_paragraph(paragraph, PALE_TEAL)
            add_inline(paragraph, " ".join(quote_lines), color=TEAL)
            continue

        numbered = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        bullet = re.match(r"^[-*]\s+(.+)$", stripped)
        checkbox = re.match(r"^-\s+\[([ xX])\]\s+(.+)$", stripped)
        if checkbox:
            flush_paragraph()
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Mm(7)
            paragraph.paragraph_format.keep_with_next = True
            add_inline(paragraph, f"□ {checkbox.group(2)}")
            index += 1
            continue
        if numbered:
            flush_paragraph()
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Mm(7)
            paragraph.paragraph_format.first_line_indent = Mm(-5)
            add_inline(paragraph, f"{numbered.group(1)}.　{numbered.group(2)}")
            index += 1
            continue
        if bullet:
            flush_paragraph()
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline(paragraph, bullet.group(1))
            index += 1
            continue

        if stripped == "---":
            flush_paragraph()
            spacer = document.add_paragraph()
            spacer.paragraph_format.space_after = Pt(2)
            index += 1
            continue

        if not stripped:
            flush_paragraph()
        else:
            paragraph_buffer.append(line)
        index += 1

    flush_paragraph()


def main() -> int:
    if not SOURCE.exists():
        print(f"Missing source: {SOURCE}", file=sys.stderr)
        return 1

    document = Document()
    configure_styles(document)
    properties = document.core_properties
    properties.title = "让 AI 看懂我的 Apple 健康"
    properties.subject = "iPhone 快捷指令 + Cloudflare + HTTP MCP 从零搭建教程"
    properties.author = "Apple Health Shortcuts MCP contributors"
    properties.keywords = "Apple Health, Shortcuts, Cloudflare, MCP, 开源教程"

    add_cover(document)

    body_section = document.add_section(WD_SECTION.CONTINUOUS)
    body_section.top_margin = Mm(17)
    body_section.bottom_margin = Mm(16)
    body_section.left_margin = Mm(17)
    body_section.right_margin = Mm(17)
    body_section.header_distance = Mm(6)
    body_section.footer_distance = Mm(7)
    add_running_header_footer(body_section)

    add_contents(document)
    add_markdown_body(document, SOURCE.read_text(encoding="utf-8"))

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
